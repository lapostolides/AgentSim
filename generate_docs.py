"""Generate AgentSim technical documentation as .docx with embedded diagrams."""

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = Path("./docs_output")
OUT_DIR.mkdir(exist_ok=True)


# ── Diagram 1: High-Level Architecture ──────────────────────────────────────

def draw_architecture_diagram():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def box(x, y, w, h, label, color="#4A90D9", fontsize=9, sublabel=None):
        rect = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.1",
            facecolor=color, edgecolor="#2C3E50", linewidth=1.2, alpha=0.9,
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2 + (0.08 if sublabel else 0), label,
                ha="center", va="center", fontsize=fontsize,
                fontweight="bold", color="white")
        if sublabel:
            ax.text(x + w / 2, y + h / 2 - 0.15, sublabel,
                    ha="center", va="center", fontsize=6.5, color="#E8E8E8")

    def arrow(x1, y1, x2, y2, label=None):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#555", lw=1.5))
        if label:
            mx, my = (x1 + x2) / 2, (y1 + y2) / 2
            ax.text(mx, my + 0.12, label, ha="center", fontsize=6.5, color="#555")

    # CLI layer
    box(3.5, 5.5, 3, 0.7, "CLI Entry Point", "#2C3E50", sublabel="main.py  (click)")

    # Orchestrator
    box(3.5, 4.0, 3, 0.7, "Orchestrator / Runner", "#8E44AD", sublabel="runner.py  •  config.py")

    # Agent Registry
    box(7.2, 4.0, 2.5, 0.7, "Agent Registry", "#27AE60", sublabel="agent_registry.py")

    # Agents row
    agents = [
        ("Lit\nScout", 0.2),
        ("Hypo-\nthesis", 1.6),
        ("Scene", 3.0),
        ("Exec-\nutor", 4.4),
        ("Eval-\nuator", 5.8),
        ("Ana-\nlyst", 7.2),
        ("Lit\nValid.", 8.6),
    ]
    colors = ["#3498DB", "#2980B9", "#2471A3", "#1F618D", "#1A5276", "#154360", "#3498DB"]
    for i, (name, x) in enumerate(agents):
        box(x, 2.2, 1.2, 0.8, name, colors[i], fontsize=7.5)

    # State layer
    box(0.5, 0.5, 4.0, 0.8, "ExperimentState", "#E74C3C", sublabel="Frozen Pydantic models  •  Tuple accumulation")

    # Transitions + Serialization
    box(5.2, 0.5, 2.2, 0.8, "Transitions", "#C0392B", sublabel="Pure functions")
    box(7.8, 0.5, 2.0, 0.8, "Serialization", "#C0392B", sublabel="JSON / Prompt ctx")

    # Environment
    box(0.2, 4.0, 2.8, 0.7, "Env Discovery", "#F39C12", sublabel="importlib probing")

    # Arrows
    arrow(5.0, 5.5, 5.0, 4.7)  # CLI -> Orchestrator
    arrow(6.5, 4.35, 7.2, 4.35)  # Orchestrator -> Registry
    arrow(5.0, 4.0, 5.0, 3.0)  # Orchestrator -> Agents
    arrow(5.0, 2.2, 5.0, 1.3)  # Agents -> State
    arrow(1.6, 4.0, 1.6, 3.0, "packages")  # Env -> Agents
    arrow(3.3, 4.0, 3.5, 4.0)  # Env -> Orchestrator
    arrow(4.5, 0.9, 5.2, 0.9)  # State -> Transitions
    arrow(7.4, 0.9, 7.8, 0.9)  # Transitions -> Serialization

    ax.set_title("AgentSim — High-Level Architecture", fontsize=13, fontweight="bold", pad=12)
    path = OUT_DIR / "architecture.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ── Diagram 2: Experiment Loop Flow ─────────────────────────────────────────

def draw_experiment_flow():
    fig, ax = plt.subplots(figsize=(9, 8.5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 9)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def phase_box(x, y, w, h, label, color, sublabel=None):
        rect = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.08",
            facecolor=color, edgecolor="#2C3E50", linewidth=1.2, alpha=0.9,
        )
        ax.add_patch(rect)
        ty = y + h / 2 + (0.07 if sublabel else 0)
        ax.text(x + w / 2, ty, label, ha="center", va="center",
                fontsize=8.5, fontweight="bold", color="white")
        if sublabel:
            ax.text(x + w / 2, y + h / 2 - 0.12, sublabel,
                    ha="center", va="center", fontsize=6.5, color="#ddd")

    def diamond(cx, cy, label):
        s = 0.45
        pts = [(cx, cy + s), (cx + s * 1.5, cy), (cx, cy - s), (cx - s * 1.5, cy)]
        poly = plt.Polygon(pts, facecolor="#F39C12", edgecolor="#2C3E50", lw=1.2)
        ax.add_patch(poly)
        ax.text(cx, cy, label, ha="center", va="center", fontsize=7, fontweight="bold", color="white")

    def arrow(x1, y1, x2, y2, label=None, color="#555"):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=1.3))
        if label:
            mx, my = (x1 + x2) / 2, (y1 + y2) / 2
            ax.text(mx + 0.15, my, label, fontsize=6.5, color=color)

    # Init phase
    phase_box(3.2, 8.2, 2.6, 0.5, "Initialize State", "#2C3E50", sublabel="start_experiment()")
    arrow(4.5, 8.2, 4.5, 7.9)

    phase_box(3.2, 7.4, 2.6, 0.5, "Discover Environment", "#F39C12", sublabel="importlib probing")
    arrow(4.5, 7.4, 4.5, 7.1)

    phase_box(3.2, 6.6, 2.6, 0.5, "Literature Scout", "#3498DB", sublabel="runs once  •  sonnet")
    arrow(4.5, 6.6, 4.5, 6.25)

    # Loop boundary
    loop_rect = FancyBboxPatch(
        (0.8, 0.6), 7.4, 5.5, boxstyle="round,pad=0.15",
        facecolor="none", edgecolor="#888", linewidth=1.5, linestyle="--",
    )
    ax.add_patch(loop_rect)
    ax.text(1.1, 5.85, "Iteration Loop  (max_iterations)", fontsize=8,
            fontweight="bold", color="#666", style="italic")

    # Phases inside loop
    phases = [
        (5.2, "Hypothesis Agent", "#2980B9", "formalize hypothesis  •  sonnet"),
        (4.5, "Scene Agent", "#2471A3", "generate Python scripts  •  sonnet"),
        (3.8, "Executor Agent", "#1F618D", "run simulations  •  sonnet"),
        (3.1, "Evaluator Agent", "#1A5276", "compute metrics  •  sonnet"),
        (2.4, "Analyst Agent", "#154360", "assess results  •  opus"),
    ]

    for y, label, color, sub in phases:
        phase_box(3.2, y, 2.6, 0.5, label, color, sublabel=sub)

    for i in range(len(phases) - 1):
        arrow(4.5, phases[i][0], 4.5, phases[i][0] - 0.2)

    # Decision diamond
    diamond(4.5, 1.65, "should\nstop?")
    arrow(4.5, 2.4, 4.5, 2.1)

    # Literature validator on the right
    phase_box(6.2, 2.4, 2.0, 0.5, "Lit Validator", "#3498DB", sublabel="sonnet")
    arrow(5.8, 2.65, 6.2, 2.65)

    # Yes -> complete
    phase_box(6.5, 1.35, 1.7, 0.45, "COMPLETED", "#27AE60")
    arrow(5.2, 1.65, 6.5, 1.55, "yes", "#27AE60")

    # No -> loop back
    ax.annotate("", xy=(2.2, 5.45), xytext=(2.2, 1.65),
                arrowprops=dict(arrowstyle="-|>", color="#E74C3C", lw=1.3,
                                connectionstyle="arc3,rad=0"))
    ax.text(1.5, 3.5, "no", fontsize=7, color="#E74C3C", fontweight="bold")

    ax.set_title("AgentSim — Experiment Loop", fontsize=13, fontweight="bold", pad=12)
    path = OUT_DIR / "experiment_flow.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ── Diagram 3: State Machine ────────────────────────────────────────────────

def draw_state_machine():
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.5)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    states = [
        (0.05, "INIT"),
        (1.2, "LIT_REV"),
        (2.35, "HYPO"),
        (3.5, "PLAN"),
        (4.65, "SCENES"),
        (5.8, "EXEC"),
        (6.95, "EVAL"),
        (8.1, "ANALYZED"),
        (9.15, "DONE"),
    ]
    colors = ["#95A5A6", "#3498DB", "#2980B9", "#2471A3", "#1F618D",
              "#1A5276", "#154360", "#8E44AD", "#27AE60"]

    for (x, label), c in zip(states, colors):
        rect = FancyBboxPatch(
            (x, 1.5), 0.95, 0.6, boxstyle="round,pad=0.06",
            facecolor=c, edgecolor="#2C3E50", linewidth=1,
        )
        ax.add_patch(rect)
        ax.text(x + 0.475, 1.8, label, ha="center", va="center",
                fontsize=6.5, fontweight="bold", color="white")

    # Arrows between states
    for i in range(len(states) - 1):
        x1 = states[i][0] + 0.95
        x2 = states[i + 1][0]
        ax.annotate("", xy=(x2, 1.8), xytext=(x1, 1.8),
                    arrowprops=dict(arrowstyle="->", color="#555", lw=1))

    # ANALYZED -> HYPO loop (continue)
    ax.annotate("", xy=(2.35, 1.5), xytext=(8.575, 1.5),
                arrowprops=dict(arrowstyle="->", color="#E74C3C", lw=1.2,
                                connectionstyle="arc3,rad=0.4"))
    ax.text(5.5, 0.45, "analyst: should_stop = false  (next iteration)", ha="center",
            fontsize=7, color="#E74C3C", style="italic")

    # FAILED state
    rect = FancyBboxPatch((4.2, 2.6), 1.2, 0.5, boxstyle="round,pad=0.06",
                          facecolor="#E74C3C", edgecolor="#2C3E50", linewidth=1)
    ax.add_patch(rect)
    ax.text(4.8, 2.85, "FAILED", ha="center", va="center",
            fontsize=7, fontweight="bold", color="white")
    ax.annotate("", xy=(4.5, 2.6), xytext=(4.0, 2.1),
                arrowprops=dict(arrowstyle="->", color="#E74C3C", lw=1, linestyle="--"))
    ax.text(3.5, 2.45, "any error", fontsize=6, color="#E74C3C")

    ax.set_title("ExperimentStatus — State Machine", fontsize=11, fontweight="bold", pad=8)
    path = OUT_DIR / "state_machine.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ── Diagram 4: Data Model Relationships ─────────────────────────────────────

def draw_data_model():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def entity(x, y, w, h, title, fields, color="#4A90D9"):
        rect = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.06",
            facecolor="white", edgecolor=color, linewidth=1.5,
        )
        ax.add_patch(rect)
        # Title bar
        title_rect = FancyBboxPatch(
            (x, y + h - 0.35), w, 0.35, boxstyle="round,pad=0.06",
            facecolor=color, edgecolor=color, linewidth=1,
        )
        ax.add_patch(title_rect)
        ax.text(x + w / 2, y + h - 0.175, title, ha="center", va="center",
                fontsize=7.5, fontweight="bold", color="white")
        for i, f in enumerate(fields):
            ax.text(x + 0.1, y + h - 0.5 - i * 0.2, f,
                    fontsize=6, color="#333", family="monospace")

    def rel_arrow(x1, y1, x2, y2, label="", style="->"):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color="#888", lw=1))
        if label:
            mx, my = (x1 + x2) / 2, (y1 + y2) / 2
            ax.text(mx, my + 0.1, label, fontsize=5.5, color="#888", ha="center")

    # ExperimentState (center top)
    entity(2.8, 4.5, 4.4, 2.3, "ExperimentState", [
        "id: str",
        "status: ExperimentStatus",
        "iteration: int",
        "raw_hypothesis: str",
        "files: tuple[FileReference]",
        "hypothesis: Hypothesis | None",
        "plan: ExperimentPlan | None",
        "scenes: tuple[SceneSpec]",
        "execution_results: tuple[ExecutionResult]",
        "evaluations: tuple[EvaluationResult]",
        "analyses: tuple[AnalysisReport]",
    ], "#E74C3C")

    # Left column
    entity(0.0, 2.5, 2.4, 1.4, "Hypothesis", [
        "formalized: str",
        "variables: list[str]",
        "parameter_space: list[ParameterSpec]",
        "predictions: list[str]",
    ], "#2980B9")

    entity(0.0, 0.5, 2.4, 1.4, "LiteratureContext", [
        "entries: tuple[LiteratureEntry]",
        "summary: str",
        "open_questions: tuple[str]",
        "methodology_notes: str",
    ], "#3498DB")

    # Center column
    entity(3.0, 2.5, 2.2, 1.2, "ExperimentPlan", [
        "hypothesis_id: str",
        "simulation_approach: str",
        "scene_descriptions: list",
        "metrics: list[str]",
    ], "#8E44AD")

    entity(3.0, 0.5, 2.2, 1.2, "SceneSpec", [
        "plan_id: str",
        "code: str  # full Python",
        "parameters: dict",
        "file_refs: list[str]",
    ], "#2471A3")

    # Right column
    entity(5.8, 2.5, 2.1, 1.0, "ExecutionResult", [
        "scene_id: str",
        "status: str",
        "stdout/stderr: str",
    ], "#1F618D")

    entity(5.8, 0.5, 2.1, 1.0, "EvaluationResult", [
        "scene_id: str",
        "metrics: dict[str,float]",
        "summary: str",
    ], "#1A5276")

    entity(8.1, 2.5, 1.8, 1.3, "AnalysisReport", [
        "findings: list[str]",
        "confidence: float",
        "supports_hypothesis",
        "should_stop: bool",
    ], "#154360")

    entity(8.1, 0.5, 1.8, 1.0, "LiteratureValidation", [
        "consistency: str",
        "novel_findings",
        "confidence_adj: float",
    ], "#F39C12")

    # Relationships
    rel_arrow(2.8, 5.0, 2.4, 3.9, "1")
    rel_arrow(3.5, 4.5, 3.5, 3.7, "1")
    rel_arrow(4.5, 4.5, 4.1, 1.7, "0..*")
    rel_arrow(5.5, 4.5, 5.8, 3.5, "0..*")
    rel_arrow(6.0, 4.5, 6.5, 3.5, "0..*")
    rel_arrow(7.2, 4.9, 8.1, 3.8, "0..*")
    rel_arrow(2.8, 4.6, 1.2, 1.9, "0..1")
    rel_arrow(7.2, 4.6, 9.0, 1.5, "0..1")

    ax.set_title("Data Model — Entity Relationships", fontsize=11, fontweight="bold", pad=8)
    path = OUT_DIR / "data_model.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


# ── Build Document ───────────────────────────────────────────────────────────

def make_doc():
    # Generate diagrams
    arch_img = draw_architecture_diagram()
    flow_img = draw_experiment_flow()
    sm_img = draw_state_machine()
    dm_img = draw_data_model()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Heading styles
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # ── Title ────────────────────────────────────────────────────────────
    title = doc.add_heading("AgentSim — Technical Documentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autonomous Hypothesis-Driven Simulation for Scientific Research")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("v0.1.0  •  Python ≥ 3.11  •  claude-agent-sdk + Pydantic")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    # ── 1. Overview ──────────────────────────────────────────────────────
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "AgentSim orchestrates seven specialized LLM agents in an iterative loop to "
        "autonomously design, execute, and evaluate scientific simulations. A user provides "
        "a natural-language hypothesis and optional data files; the system formalizes the "
        "hypothesis, generates runnable simulation code, executes it, evaluates results, "
        "and decides whether to iterate or conclude — all without human intervention."
    )

    # Key properties
    doc.add_heading("Core Design Properties", level=2)
    props = [
        ("Immutable state", "All models are frozen Pydantic objects. Transitions return new instances; nothing is mutated."),
        ("Domain-agnostic", "Not tied to any simulation domain. Agents discover available packages at runtime."),
        ("Self-terminating", "The Analyst agent decides when confidence is sufficient or diminishing returns are reached."),
        ("Literature-grounded", "A literature survey runs first and its context flows into every subsequent agent prompt."),
    ]
    for title_text, desc in props:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    # ── 2. Architecture ──────────────────────────────────────────────────
    doc.add_heading("2. Architecture", level=1)
    doc.add_picture(str(arch_img), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Layer table
    doc.add_heading("Layer Summary", level=2)
    table = doc.add_table(rows=5, cols=3, style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Layer", "Modules", "Responsibility"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    rows = [
        ("CLI", "main.py", "Click commands: run, interactive"),
        ("Orchestrator", "runner.py, config.py", "Sequences phases, manages budget, persists state"),
        ("Agents", "7 agent modules + registry", "LLM-backed phases with typed JSON output"),
        ("State", "models.py, transitions.py,\nserialization.py", "Frozen models, pure transitions, prompt serialization"),
    ]
    for i, (layer, mods, resp) in enumerate(rows):
        table.rows[i + 1].cells[0].text = layer
        table.rows[i + 1].cells[1].text = mods
        table.rows[i + 1].cells[2].text = resp

    # ── 3. Experiment Flow ───────────────────────────────────────────────
    doc.add_heading("3. Experiment Lifecycle", level=1)
    doc.add_picture(str(flow_img), width=Inches(5.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Phase Sequence", level=2)
    doc.add_paragraph(
        "Each iteration runs six agent phases in strict order. "
        "The Literature Scout runs once before the loop. "
        "Each phase is a separate query() call to the claude-agent-sdk, "
        "keeping context windows clean."
    )

    # Agent table
    table = doc.add_table(rows=8, cols=4, style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["#", "Agent", "Model", "Output"]):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    agent_rows = [
        ("0", "Literature Scout", "sonnet", "LiteratureContext (papers, methodology notes)"),
        ("1", "Hypothesis", "sonnet", "Structured Hypothesis (variables, params, predictions)"),
        ("2", "Scene", "sonnet", "SceneSpec[] — complete runnable Python scripts"),
        ("3", "Executor", "sonnet", "ExecutionResult[] — stdout, stderr, output paths"),
        ("4", "Evaluator", "sonnet", "EvaluationResult[] — metrics, artifacts"),
        ("5", "Analyst", "opus", "AnalysisReport — confidence, should_stop decision"),
        ("6", "Literature Validator", "sonnet", "LiteratureValidation — confidence adjustment"),
    ]
    for i, (num, name, model, output) in enumerate(agent_rows):
        table.rows[i + 1].cells[0].text = num
        table.rows[i + 1].cells[1].text = name
        table.rows[i + 1].cells[2].text = model
        table.rows[i + 1].cells[3].text = output

    # ── 4. State Machine ─────────────────────────────────────────────────
    doc.add_heading("4. State Machine", level=1)
    doc.add_picture(str(sm_img), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "State transitions are pure functions in transitions.py: "
        "fn(state, data) → new_state via model_copy(update={...}). "
        "Results accumulate in tuples — scenes, execution_results, evaluations, "
        "and analyses are appended, never replaced. This preserves full experiment history."
    )

    doc.add_heading("Key Transition Behaviors", level=2)
    behaviors = [
        "add_analysis() has dual responsibility: increments iteration counter AND sets status to COMPLETED (should_stop=true) or ANALYZED (should_stop=false).",
        "start_experiment() auto-detects file types from extensions (e.g., .stl → mesh, .yaml → config).",
        "State >50KB is truncated for prompts: counts replace full lists, only latest analysis is kept.",
    ]
    for b in behaviors:
        doc.add_paragraph(b, style="List Bullet")

    # ── 5. Data Model ────────────────────────────────────────────────────
    doc.add_heading("5. Data Model", level=1)
    doc.add_picture(str(dm_img), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "ExperimentState is the single top-level envelope. "
        "All nested models are frozen (immutable). "
        "Multi-valued fields use tuple[T, ...] for immutability. "
        "Each agent phase reads serialized state and returns structured JSON "
        "that is validated into the appropriate Pydantic model."
    )

    # ── 6. Environment Discovery ─────────────────────────────────────────
    doc.add_heading("6. Environment Discovery", level=1)
    doc.add_paragraph(
        "At startup, discover_environment() probes for known packages "
        "(numpy, scipy, mitsuba, opencv, matplotlib, etc.) using "
        "importlib.util.find_spec — no imports executed. "
        "Results are formatted into a human-readable string "
        "baked into agent prompts so they know exactly which packages "
        "and versions are available when generating simulation code."
    )
    doc.add_paragraph(
        "Custom packages can be added via OrchestratorConfig.extra_packages "
        "(dict mapping display_name → import_name)."
    )

    # ── 7. JSON Extraction ───────────────────────────────────────────────
    doc.add_heading("7. JSON Extraction Strategy", level=1)
    doc.add_paragraph(
        "Agent responses may contain prose around JSON. "
        "The runner uses a four-tier fallback:"
    )
    tiers = [
        "Direct JSON.parse of full response text",
        "Extract from markdown code fences (```json ... ```)",
        "Find content between first { and last }",
        "Phase-specific fallback (e.g., Analyst defaults to should_stop=true)",
    ]
    for i, t in enumerate(tiers):
        doc.add_paragraph(f"{i+1}. {t}")

    # ── 8. Configuration ─────────────────────────────────────────────────
    doc.add_heading("8. Configuration", level=1)

    table = doc.add_table(rows=7, cols=3, style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Parameter", "Default", "Description"]):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    config_rows = [
        ("max_iterations", "5", "Maximum experiment loop iterations"),
        ("max_budget_usd", "10.0", "Total budget across all phases"),
        ("max_turns_per_phase", "30", "Max agent turns per phase"),
        ("output_dir", "./output", "Directory for state and artifacts"),
        ("save_intermediate_state", "True", "Persist final_state.json"),
        ("extra_packages", "{}", "Additional packages for discovery"),
    ]
    for i, (param, default, desc) in enumerate(config_rows):
        table.rows[i + 1].cells[0].text = param
        table.rows[i + 1].cells[1].text = default
        table.rows[i + 1].cells[2].text = desc

    # ── 9. Quick Reference ───────────────────────────────────────────────
    doc.add_heading("9. Quick Reference", level=1)

    doc.add_heading("Commands", level=2)
    cmds = [
        ("pip install -e '.[dev]'", "Install with dev dependencies"),
        ("pytest", "Run all tests"),
        ("pytest tests/unit/", "Unit tests only"),
        ("pytest tests/unit/test_file.py::test_name", "Single test"),
        ("ruff check src/ tests/", "Lint"),
        ("agentsim run \"hypothesis\"", "Run experiment"),
        ("agentsim interactive", "Interactive REPL"),
    ]
    table = doc.add_table(rows=len(cmds) + 1, cols=2, style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Command"
    table.rows[0].cells[1].text = "Purpose"
    for p in table.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in table.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.bold = True
    for i, (cmd, purpose) in enumerate(cmds):
        table.rows[i + 1].cells[0].text = cmd
        table.rows[i + 1].cells[1].text = purpose

    doc.add_heading("Dependencies", level=2)
    doc.add_paragraph(
        "Core: claude-agent-sdk, pydantic ≥2.0, anyio, click, structlog. "
        "Dev: pytest, pytest-asyncio (asyncio_mode=auto), ruff, fastmcp."
    )

    # ── Save ─────────────────────────────────────────────────────────────
    output_path = Path("./AgentSim_Documentation.docx")
    doc.save(output_path)
    print(f"Documentation saved to {output_path}")
    return output_path


if __name__ == "__main__":
    make_doc()
